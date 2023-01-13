import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx.shared import Pt
from docx import Document
import datetime

class WebScraper:
    def __init__(self, website):
        self.website = website

    def get_soup(self, url):
        r = requests.get(url)
        soup = BeautifulSoup(r.text, 'html.parser')
        return soup

    def get_href(self, soup):
        href = []
        for link in soup.find_all('a', class_='DTTitulo'):
            href.append(link.get('href'))
        return href

    def get_url(self, links):
        url = []
        for link in links:
            url.append('https://es.catholic.net' + link)
        return url

    def get_title(self, soup):
        title = []
        for link in soup.find_all('a', class_='DTTitulo'):
            title.append(link.get_text())
        return title

    def get_texto(self, soup):
        texto = []
        for link in soup.find_all('div', id='DTTexto'):
            texto.append(link.get_text())
        return texto

    def get_texto_from_url(self, url):
        texto = []
        for link in url:
            soup = self.get_soup(link)
            for i in soup.find('div', {'id': 'art_texto'}).text:
                texto.append(i)
        return texto

    def export_df_to_doc_v3(self, df, file_path):
        document = Document()

        # Agregar encabezado
        header = document.sections[0].header
        header.paragraphs[0].text = 'Un lugar de encuentro para católicos'
        header.paragraphs[0].style = 'Header'

        # Agregar pie de página
        footer = document.sections[0].footer
        footer.paragraphs[0].text = 'Realizado por @MundoBits'
        footer.paragraphs[0].style = 'Footer'

        # Agregar contenido del DataFrame al documento
        for row in df.itertuples():
            p = document.add_paragraph("")
            run = p.add_run(row.Tema)
            run.bold = True
            run.font.size = Pt(18)
            p = document.add_paragraph("")
            run = p.add_run(row.Titulo)
            run.bold = True
            run.font.size = Pt(16)
            p = document.add_paragraph("")
            run = p.add_run(str(row.Entrada))
            run.italic = True
            run.font.size = Pt(14)
            document.add_paragraph("" + row.Texto)
            document.add_paragraph(" ")

        # Guardar documento
        document.save(file_path)

    def main(self):
        soup = self.get_soup(self.website)
        links = self.get_href(soup)
        url = self.get_url(links)
        title = self.get_title(soup)
        texto = self.get_texto(soup)

        text_art = []

        for i in range(3):
            try:
                response = requests.get(url[i])
                response.raise_for_status()

                soup = BeautifulSoup(response.text, 'html.parser')
                text_art.append(soup.find('div', {'id': 'art_texto'}).text)

            except requests.exceptions.ConnectionError as e:
                print(e)

        temas = ['El Evangelio meditado', 'Meditacion para hoy', 'Tema actual']
        df = pd.DataFrame({'Tema': temas, 'Titulo': title[:3], 'Entrada': texto[:3], 'Texto': text_art})

        # llamar la funcion export_df_to_doc_v3 y guardar el archivo con catholic + fecha de hoy
        fecha = datetime.datetime.now().strftime("%Y-%m-%d")
        self.export_df_to_doc_v3(df, 'catholic' + fecha + '.docx')

if __name__ == '__main__':
    website = 'https://es.catholic.net/'
    scraper = WebScraper(website)
    scraper.main()
